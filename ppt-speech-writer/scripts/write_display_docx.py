#!/usr/bin/env python3
"""Write a complete display-version speaker-notes document.

Input is display_document.json. When python-docx is installed, the output is a
.docx file. Otherwise a Markdown fallback is written next to the requested path.
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path


def as_items(value) -> list[tuple[str, str]]:
    if isinstance(value, dict):
        return [(str(k), str(v)) for k, v in value.items()]
    if isinstance(value, list):
        return [(str(i + 1), str(v)) for i, v in enumerate(value)]
    if value:
        return [("", str(value))]
    return []


def markdown_table(headers: list[str], rows: list[list[str]]) -> list[str]:
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        clean = [str(cell).replace("\n", " ").replace("|", "/") for cell in row]
        lines.append("| " + " | ".join(clean) + " |")
    return lines


TIMING_HEADERS = ["Slide", "Title", "Est. time", "Words/Chars", "Budget", "Pauses"]


def _to_int(value) -> int | None:
    try:
        return int(float(value))
    except (TypeError, ValueError):
        return None


def _parse_time(value) -> int | None:
    text = str(value or "").strip()
    if not text or ":" not in text:
        return None
    parts = text.split(":")
    try:
        return int(parts[0]) * 60 + int(parts[1])
    except ValueError:
        return None


def _fmt_time(seconds: int) -> str:
    minutes, secs = divmod(int(seconds), 60)
    return f"{minutes}:{secs:02d}"


def timing_rows(items: list[dict]) -> tuple[list[list[str]], list[str]]:
    """Build timing rows plus a TOTAL row summing counts, budgets, pauses, and time."""
    rows: list[list[str]] = []
    total_count = total_budget = total_pauses = total_seconds = 0
    have_time = False
    for item in items:
        rows.append(
            [
                item.get("slide", ""),
                item.get("title", ""),
                item.get("time", ""),
                item.get("word_count", ""),
                item.get("budget", ""),
                item.get("pauses", ""),
            ]
        )
        total_count += _to_int(item.get("word_count")) or 0
        total_budget += _to_int(item.get("budget")) or 0
        total_pauses += _to_int(item.get("pauses")) or 0
        seconds = _parse_time(item.get("time"))
        if seconds is not None:
            total_seconds += seconds
            have_time = True
    total_row = [
        "TOTAL",
        "",
        _fmt_time(total_seconds) if have_time else "",
        str(total_count) if total_count else "",
        str(total_budget) if total_budget else "",
        str(total_pauses) if total_pauses else "",
    ]
    return rows, total_row


def write_markdown(data: dict, output: Path) -> Path:
    md_path = output.with_suffix(".md")
    md_path.parent.mkdir(parents=True, exist_ok=True)
    lines: list[str] = [
        f"# {data.get('title', 'Speaker Notes Display Version')}",
        "",
        f"Deck: {data.get('deck_path', '')}",
        "",
        "## Deck Comprehension Brief",
        "",
    ]
    for key, value in as_items(data.get("comprehension_brief", {})):
        lines.extend([f"### {key}", str(value), ""])

    lines.extend(["## Narrative Arc", ""])
    for key, value in as_items(data.get("narrative_arc", {})):
        lines.append(f"- **{key}:** {value}")
    lines.append("")

    lines.extend(["## Slide-by-Slide Display Notes", ""])
    for slide in data.get("slides", []):
        lines.extend(
            [
                f"### Slide {slide.get('slide')} - {slide.get('title', '')}",
                "",
                str(slide.get("display_notes", "")).strip(),
                "",
            ]
        )

    lines.extend(["## Key Parameters And Methods", ""])
    kpm_rows = [
        [item.get("term", ""), item.get("type", ""), item.get("slides", ""), item.get("definition", "")]
        for item in data.get("key_parameters_methods", [])
    ]
    lines.extend(markdown_table(["Term", "Type", "Slide(s)", "Definition"], kpm_rows))
    lines.append("")

    lines.extend(["## Timing Table", ""])
    rows, total_row = timing_rows(data.get("timing", []))
    lines.extend(markdown_table(TIMING_HEADERS, rows + [total_row]))
    lines.append("")

    lines.extend(["## Coverage Notes", ""])
    for note in data.get("coverage_notes", []):
        lines.append(f"- {note}")
    lines.append("")

    lines.extend(["## Injection Log", ""])
    for line in data.get("injection_log", []):
        lines.append(f"- {line}")
    lines.append("")

    md_path.write_text("\n".join(lines), encoding="utf-8")
    return md_path


def write_docx(data: dict, output: Path) -> Path:
    try:
        from docx import Document
    except ImportError:
        return write_markdown(data, output)

    doc = Document()
    doc.add_heading(data.get("title", "Speaker Notes Display Version"), level=0)
    if data.get("deck_path"):
        doc.add_paragraph(f"Deck: {data['deck_path']}")

    doc.add_heading("Deck Comprehension Brief", level=1)
    for key, value in as_items(data.get("comprehension_brief", {})):
        doc.add_heading(key, level=2)
        doc.add_paragraph(value)

    doc.add_heading("Narrative Arc", level=1)
    for key, value in as_items(data.get("narrative_arc", {})):
        doc.add_paragraph(f"{key}: {value}", style="List Bullet")

    doc.add_heading("Slide-by-Slide Display Notes", level=1)
    for slide in data.get("slides", []):
        doc.add_heading(f"Slide {slide.get('slide')} - {slide.get('title', '')}", level=2)
        for paragraph in str(slide.get("display_notes", "")).split("\n"):
            doc.add_paragraph(paragraph)

    doc.add_heading("Key Parameters And Methods", level=1)
    table = doc.add_table(rows=1, cols=4)
    for idx, header in enumerate(["Term", "Type", "Slide(s)", "Definition"]):
        table.rows[0].cells[idx].text = header
    for item in data.get("key_parameters_methods", []):
        row = table.add_row().cells
        row[0].text = str(item.get("term", ""))
        row[1].text = str(item.get("type", ""))
        row[2].text = str(item.get("slides", ""))
        row[3].text = str(item.get("definition", ""))

    doc.add_heading("Timing Table", level=1)
    timing = doc.add_table(rows=1, cols=len(TIMING_HEADERS))
    for idx, header in enumerate(TIMING_HEADERS):
        timing.rows[0].cells[idx].text = header
    rows, total_row = timing_rows(data.get("timing", []))
    for values in rows + [total_row]:
        cells = timing.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = str(value)

    doc.add_heading("Coverage Notes", level=1)
    for note in data.get("coverage_notes", []):
        doc.add_paragraph(str(note), style="List Bullet")

    doc.add_heading("Injection Log", level=1)
    for line in data.get("injection_log", []):
        doc.add_paragraph(str(line), style="List Bullet")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser(description="Write a complete display-version speaker-notes document.")
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args(argv)

    if not args.input.exists():
        sys.stderr.write(f"Input not found: {args.input}\n")
        return 1
    data = json.loads(args.input.read_text(encoding="utf-8"))
    written = write_docx(data, args.output)
    print(f"saved: {written}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
